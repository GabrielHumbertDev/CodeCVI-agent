"""
Export service: convert tailored CV data (dict) to DOCX or PDF bytes.
"""
import io
from typing import Optional

# ── DOCX ────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── PDF ─────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _contact_line(data: dict) -> str:
    parts = [p for p in [data.get("email"), data.get("phone")] if p]
    return " | ".join(parts)


def _section_title_docx(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text.upper()).bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    # Underline via border trick
    border = p._p.get_or_add_pPr()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_to_docx(data: dict) -> bytes:
    doc = Document()

    # ── Margins ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Name ─────────────────────────────────────────────────────────────────
    name = data.get("name", "")
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Contact ───────────────────────────────────────────────────────────────
    contact = _contact_line(data)
    if contact:
        c_para = doc.add_paragraph(contact)
        c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()  # spacer

    # ── Summary ───────────────────────────────────────────────────────────────
    summary = data.get("summary", "")
    if summary:
        _add_heading(doc, "Professional Summary")
        doc.add_paragraph(summary)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = data.get("skills", [])
    if skills:
        _add_heading(doc, "Skills")
        doc.add_paragraph(" • ".join(skills))

    # ── Experience ────────────────────────────────────────────────────────────
    experience = data.get("experience", [])
    if experience:
        _add_heading(doc, "Experience")
        for exp in experience:
            title = exp.get("title", "")
            company = exp.get("company", "")
            dates = exp.get("dates", "")
            description = exp.get("description", "")

            p = doc.add_paragraph()
            run = p.add_run(f"{title}")
            run.bold = True
            if company:
                p.add_run(f"  —  {company}")
            if dates:
                p.add_run(f"  ({dates})")
            p.paragraph_format.space_before = Pt(4)

            if description:
                desc_para = doc.add_paragraph(description)
                desc_para.paragraph_format.left_indent = Pt(12)
                desc_para.paragraph_format.space_after = Pt(4)

            for bullet in exp.get("bullets", []):
                b = doc.add_paragraph(style="List Bullet")
                b.add_run(bullet)
                b.paragraph_format.left_indent = Pt(12)

    # ── Education ─────────────────────────────────────────────────────────────
    education = data.get("education", [])
    if education:
        _add_heading(doc, "Education")
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            dates = edu.get("dates", "")

            p = doc.add_paragraph()
            run = p.add_run(degree)
            run.bold = True
            if institution:
                p.add_run(f"  —  {institution}")
            if dates:
                p.add_run(f"  ({dates})")
            p.paragraph_format.space_before = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def export_to_pdf(data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    base_styles = getSampleStyleSheet()

    style_name = ParagraphStyle(
        "CVName",
        parent=base_styles["Title"],
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=2,
    )
    style_contact = ParagraphStyle(
        "CVContact",
        parent=base_styles["Normal"],
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
        spaceAfter=8,
    )
    style_section = ParagraphStyle(
        "CVSection",
        parent=base_styles["Normal"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1a1a2e"),
        fontName="Helvetica-Bold",
        spaceBefore=10,
        spaceAfter=3,
    )
    style_body = ParagraphStyle(
        "CVBody",
        parent=base_styles["Normal"],
        fontSize=9,
        leading=13,
        spaceAfter=3,
    )
    style_entry_title = ParagraphStyle(
        "CVEntryTitle",
        parent=base_styles["Normal"],
        fontSize=10,
        leading=13,
        fontName="Helvetica-Bold",
        spaceBefore=5,
        spaceAfter=1,
    )
    style_bullet = ParagraphStyle(
        "CVBullet",
        parent=base_styles["Normal"],
        fontSize=9,
        leading=13,
        leftIndent=12,
        spaceAfter=1,
    )

    story = []

    # ── Name ─────────────────────────────────────────────────────────────────
    name = data.get("name", "")
    story.append(Paragraph(name, style_name))

    # ── Contact ───────────────────────────────────────────────────────────────
    contact = _contact_line(data)
    if contact:
        story.append(Paragraph(contact, style_contact))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Spacer(1, 4 * mm))

    # ── Summary ───────────────────────────────────────────────────────────────
    summary = data.get("summary", "")
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", style_section))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd")))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(summary, style_body))

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = data.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", style_section))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd")))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(" • ".join(skills), style_body))

    # ── Experience ────────────────────────────────────────────────────────────
    experience = data.get("experience", [])
    if experience:
        story.append(Paragraph("EXPERIENCE", style_section))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd")))
        story.append(Spacer(1, 2 * mm))
        for exp in experience:
            title = exp.get("title", "")
            company = exp.get("company", "")
            dates = exp.get("dates", "")
            description = exp.get("description", "")

            entry_line = title
            if company:
                entry_line += f" &mdash; {company}"
            if dates:
                entry_line += f" <font color='#777777'>({dates})</font>"

            story.append(Paragraph(entry_line, style_entry_title))
            if description:
                story.append(Paragraph(description, style_body))
            for bullet in exp.get("bullets", []):
                story.append(Paragraph(f"• {bullet}", style_bullet))

    # ── Education ─────────────────────────────────────────────────────────────
    education = data.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", style_section))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd")))
        story.append(Spacer(1, 2 * mm))
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            dates = edu.get("dates", "")

            entry_line = degree
            if institution:
                entry_line += f" &mdash; {institution}"
            if dates:
                entry_line += f" <font color='#777777'>({dates})</font>"

            story.append(Paragraph(entry_line, style_entry_title))

    doc.build(story)
    return buf.getvalue()
